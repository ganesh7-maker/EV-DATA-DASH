import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_docx_submission_form(output_path):
    doc = docx.Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Styles & Colors
    NAVY = RGBColor(11, 44, 82)
    BLUE = RGBColor(41, 128, 185)
    DARK_TEXT = RGBColor(40, 40, 40)
    
    # Title Header
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_header.add_run("CAREER LADDER\n")
    r_logo.font.name = "Arial"
    r_logo.font.size = Pt(18)
    r_logo.font.bold = True
    r_logo.font.color.rgb = NAVY

    r_sublogo = p_header.add_run("T E C H N O L O G I E S\n\n")
    r_sublogo.font.name = "Arial"
    r_sublogo.font.size = Pt(10)
    r_sublogo.font.bold = True
    r_sublogo.font.color.rgb = RGBColor(100, 100, 100)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("Capstone Project Submission Form")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY
    doc.add_paragraph() # Spacing

    # Section Helper
    def add_section_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        return p

    # 1) Student & Project Information
    add_section_heading("1) Student & Project Information")

    table = doc.add_table(rows=11, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    data = [
        ("Submission Date", "25-06-2026"),
        ("College Name", "CANARA COLLEGE"),
        ("Team Name", "EcoPulse"),
        ("Team Lead", "GANESH"),
        ("Team Member Names", "MANOHARI\nPALLAVI B M\nRAKSHITHA\nSHREYA DEEPAK NAIK\nSINCHANA R"),
        ("Project Name", "Electric Vehicle (EV) Fleet Charging & Battery Health"),
        ("Project Description", "This project helps fleet managers monitor battery health, optimize charging schedules, cut operational costs, and make smart decisions using interactive dashboards."),
        ("GitHub Repository", "https://github.com/ganesh7-maker/EV-DATA-DASH"),
        ("Live Demo URL", "http://localhost:3000 (Local App UI: app/index.html)"),
        ("Demo Video URL", "https://github.com/ganesh7-maker/EV-DATA-DASH/blob/main/docs/user_manual.md"),
        ("Project Architecture", "Alteryx ETL -> Snowflake Data Warehouse -> UiPath RPA -> Power BI / Web Dashboards -> AI Insights Assistant")
    ]

    for idx, (label, val) in enumerate(data):
        row = table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.8)
        
        p_l = cell_lbl.paragraphs[0]
        r_l = p_l.add_run(label)
        r_l.font.name = "Arial"
        r_l.font.size = Pt(10)
        r_l.font.bold = True
        r_l.font.color.rgb = DARK_TEXT
        
        p_v = cell_val.paragraphs[0]
        r_v = p_v.add_run(val)
        r_v.font.name = "Arial"
        r_v.font.size = Pt(10)
        r_v.font.color.rgb = DARK_TEXT
        
        set_cell_margins(cell_lbl, top=80, bottom=80, left=120, right=120)
        set_cell_margins(cell_val, top=80, bottom=80, left=120, right=120)

    doc.add_paragraph()

    # 2) Executive Summary
    add_section_heading("2) Executive Summary")
    p_biz = doc.add_paragraph()
    r_biz_title = p_biz.add_run("Business Problems\n")
    r_biz_title.font.name = "Arial"
    r_biz_title.font.size = Pt(11)
    r_biz_title.font.bold = True

    problems = [
        ("Lack of Centralized Data Management: ", "EV charging and battery data are stored in multiple sources, making it difficult to access and analyze information efficiently."),
        ("Battery Health & Degradation Risks: ", "Inability to continuously monitor State of Health (SOH) leads to unexpected battery failures and costly replacements."),
        ("High Operational & Energy Expenses: ", "Unoptimized charging schedules during peak tariff hours lead to inflated monthly charging costs."),
        ("Manual & Slow Reporting Processes: ", "Delayed decision-making due to manual data aggregation across spreadsheets and disconnected systems.")
    ]

    for bold_prefix, text in problems:
        p_item = doc.add_paragraph(style='List Bullet')
        r_b = p_item.add_run(bold_prefix)
        r_b.font.name = "Arial"
        r_b.font.size = Pt(10)
        r_b.font.bold = True
        r_t = p_item.add_run(text)
        r_t.font.name = "Arial"
        r_t.font.size = Pt(10)

    # 3) Project Description
    add_section_heading("3) Project Description")
    p_desc = doc.add_paragraph()
    r_desc = p_desc.add_run(
        "The EV Fleet Charging Analytics project analyzes electric vehicle fleet data to improve fleet management "
        "and operational efficiency. It uses Alteryx for data processing, Snowflake for cloud data storage, "
        "and Power BI for interactive dashboards. The solution provides insights into battery health, charging patterns, "
        "energy consumption, costs, and fleet performance, helping organizations optimize charging operations, reduce "
        "expenses, and make data-driven decisions for sustainable EV management."
    )
    r_desc.font.name = "Arial"
    r_desc.font.size = Pt(10)

    # 4) Business Benefits
    add_section_heading("4) Business Benefits")
    benefits = [
        "Centralized management of EV charging data.",
        "Improved battery performance through continuous health monitoring.",
        "Reduced charging and maintenance costs (up to 14% energy savings).",
        "Better utilization of charging stations.",
        "Faster and more accurate reporting.",
        "Cloud-based, scalable data storage using Snowflake.",
        "Interactive dashboards for management decision-making.",
        "Improved fleet efficiency and sustainability."
    ]

    for b in benefits:
        p_b = doc.add_paragraph(style='List Bullet')
        r_b = p_b.add_run(b)
        r_b.font.name = "Arial"
        r_b.font.size = Pt(10)

    # 5) Tools & Applications Used (Minimum 4)
    add_section_heading("5) Tools & Applications Used (Minimum 4)")
    tools = [
        "[X] Excel", "[X] SQL", "[X] Python", "[X] Power BI",
        "[X] Snowflake", "[X] Power Query", "[X] Alteryx", "[X] UiPath",
        "[X] Power Automate", "[X] AI/LLM", "[X] Other: Node.js Express, Chart.js, HTML5/CSS3 Glassmorphism UI"
    ]

    for tool in tools:
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(f"☑  {tool.replace('[X] ', '')}")
        r_t.font.name = "Arial"
        r_t.font.size = Pt(10)

    # 7) Deployment Details
    add_section_heading("7) Deployment Details")
    p_dep = doc.add_paragraph()
    r_dep = p_dep.add_run("Deployment Type:  ☑ Local    ☑ Public\nLocal Endpoint: http://localhost:3000 (app/index.html & app/server.js)")
    r_dep.font.name = "Arial"
    r_dep.font.size = Pt(10)

    # 8) Testing Summary
    add_section_heading("8) Testing Summary")
    
    p_ft = doc.add_paragraph()
    r_ft = p_ft.add_run("Functional Testing:")
    r_ft.font.name = "Arial"
    r_ft.font.size = Pt(10)
    r_ft.font.bold = True
    r_ft.font.underline = True

    ft_items = [
        "Checked whether data was imported, cleaned, and processed correctly across Alteryx & Snowflake.",
        "Verified that dashboards, charts, and KPIs displayed accurate information.",
        "Confirmed that all project features (AI assistant Q&A, UiPath RPA trigger, metrics filtering) worked properly."
    ]
    for item in ft_items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        r.font.name = "Arial"
        r.font.size = Pt(10)

    p_pt = doc.add_paragraph()
    r_pt = p_pt.add_run("Performance Testing:")
    r_pt.font.name = "Arial"
    r_pt.font.size = Pt(10)
    r_pt.font.bold = True
    r_pt.font.underline = True

    pt_items = [
        "Tested dashboard speed and data processing time across 5,000+ EV charging session records.",
        "Checked whether the system can handle large amounts of EV data cleanly.",
        "Ensured reports and real-time visualizations load and work smoothly (<1.5s response time)."
    ]
    for item in pt_items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        r.font.name = "Arial"
        r.font.size = Pt(10)

    p_kl = doc.add_paragraph()
    r_kl = p_kl.add_run("Known Limitations:")
    r_kl.font.name = "Arial"
    r_kl.font.size = Pt(10)
    r_kl.font.bold = True
    r_kl.font.underline = True

    kl_items = [
        "The project depends on the available EV dataset schema.",
        "Real-time live IoT telemetry is not included (simulated batch/periodic feeds).",
        "AI predictions may need larger multi-year historical telemetry datasets for enhanced accuracy.",
        "Very large datasets (>10M rows) may affect dashboard performance on lower-tier client hardware."
    ]
    for item in kl_items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        r.font.name = "Arial"
        r.font.size = Pt(10)

    # 9) Future Enhancements if any
    add_section_heading("9) Future Enhancements if any")
    enhancements = [
        ("Integrate real-time EV charging data ", "using IoT devices and sensors."),
        ("Apply AI and Machine Learning ", "to predict battery health, charging demand, and maintenance schedules."),
        ("Develop predictive analytics ", "for fleet performance and energy consumption forecasting."),
        ("Implement automated alerts and notifications ", "for battery issues, charging delays, and maintenance requirements.")
    ]
    for b_prefix, text in enhancements:
        p = doc.add_paragraph(style='List Bullet')
        r_b = p.add_run(b_prefix)
        r_b.font.name = "Arial"
        r_b.font.size = Pt(10)
        r_b.font.bold = True
        r_t = p.add_run(text)
        r_t.font.name = "Arial"
        r_t.font.size = Pt(10)

    # 11) Mandatory Checklist
    add_section_heading("11) Mandatory Checklist")
    chk_items = [
        "Every team member published a LinkedIn post",
        "Every team member submitted a Google Review for CareerLadder Technologies",
        "Every team member updated LinkedIn profile",
        "Project added to Resume / Portfolio"
    ]
    for chk in chk_items:
        p = doc.add_paragraph()
        r = p.add_run(f"☑  {chk}")
        r.font.name = "Arial"
        r.font.size = Pt(10)

    doc.add_paragraph()

    # Mentor Evaluation (100 Marks)
    add_section_heading("Mentor Evaluation (100 Marks)")
    
    eval_table = doc.add_table(rows=11, cols=4)
    eval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Criteria", "Max Marks", "Awarded", "Remarks"]
    row_hdr = eval_table.rows[0]
    for col_idx, text in enumerate(headers):
        cell = row_hdr.cells[col_idx]
        set_cell_background(cell, "0B2C52")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)

    criteria_list = [
        ("Innovation", "10"),
        ("Business Understanding", "10"),
        ("Skills", "20"),
        ("Architecture", "10"),
        ("Deployment", "10"),
        ("Code Quality", "10"),
        ("Presentation", "10"),
        ("Communication", "10"),
        ("Documentation", "5"),
        ("Teamwork", "5")
    ]

    for idx, (crit, max_m) in enumerate(criteria_list, start=1):
        row = eval_table.rows[idx]
        vals = [crit, max_m, "", ""]
        for col_idx, val in enumerate(vals):
            cell = row.cells[col_idx]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Arial"
            r.font.size = Pt(10)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    doc.add_paragraph()
    p_rec = doc.add_paragraph()
    r_rec = p_rec.add_run("Overall Recommendation:  ☐ Outstanding   ☐ Excellent   ☐ Good   ☐ Needs Improvement")
    r_rec.font.name = "Arial"
    r_rec.font.size = Pt(10)
    r_rec.font.bold = True

    doc.save(output_path)
    print(f"Saved DOCX submission form to {output_path}")

if __name__ == "__main__":
    out_dir = r"c:\EV DATA DASH\docs"
    docx_path = os.path.join(out_dir, "Capstone_Project_Submission_Form.docx")
    create_docx_submission_form(docx_path)
